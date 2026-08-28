from apps.word_doc_services.creating_document import *
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def test_creating_empty_doc():
    wrapper = BusinessCaseWordDocumentWrapper()
    assert wrapper.doc is not None


def test_creating_paragraphs():
    # Arrange
    wrapper = BusinessCaseWordDocumentWrapper()

    # Act
    wrapper.add_paragraph("This is a test")

    # Assert
    assert len(wrapper.doc.paragraphs) == 1
    assert wrapper.doc.paragraphs[0].text == "This is a test"


def test_creating_h1_section():
    # Arrange
    wrapper = BusinessCaseWordDocumentWrapper()

    # Act
    wrapper.add_h1_section_header("Test Header")

    # Assert
    assert len(wrapper.doc.paragraphs) == 1
    assert wrapper.doc.paragraphs[0].text == "Test Header"
    assert wrapper.doc.paragraphs[0].runs[0].bold == True
    assert wrapper.doc.paragraphs[0].runs[0].font.size == Pt(28)
    assert wrapper.doc.paragraphs[0].runs[0].font.name == bold_font_name


def test_creating_h2_section():
    # Arrange
    wrapper = BusinessCaseWordDocumentWrapper()

    # Act
    wrapper.add_h2_section_header("Test Header")

    # Assert
    assert len(wrapper.doc.paragraphs) == 1
    assert wrapper.doc.paragraphs[0].text == "Test Header"
    assert wrapper.doc.paragraphs[0].runs[0].bold == True
    assert wrapper.doc.paragraphs[0].runs[0].font.size == Pt(16)
    assert wrapper.doc.paragraphs[0].runs[0].font.name == bold_font_name


def test_adding_hyperlink():
    # Arrange
    wrapper = BusinessCaseWordDocumentWrapper()
    link_target: str = ""

    # Act
    wrapper.add_hyperlink("SomeLink.com", "This is a link", "link")

    rels = wrapper.doc.part.rels
    for rel in rels:
        if rels[rel].reltype == RT.HYPERLINK:
            link_target = f"{rels[rel]._target}"

    # Assert
    assert link_target == "SomeLink.com"


def test_adding_bullet_point_list():
    # Arrange
    list_of_bullet_items = ["one", "two", "three"]
    wrapper = BusinessCaseWordDocumentWrapper()

    # Act
    wrapper.add_bullet_point_list(list_of_bullet_items)

    # Assert
    assert len(wrapper.doc.paragraphs) == len(list_of_bullet_items)
    assert wrapper.doc.paragraphs[0].text == list_of_bullet_items[0]
    assert wrapper.doc.paragraphs[1].text == list_of_bullet_items[1]
    assert wrapper.doc.paragraphs[2].text == list_of_bullet_items[2]


def test_adding_fixed_text():
    # Arrange
    wrapper = BusinessCaseWordDocumentWrapper()

    # Act
    wrapper.add_fixed_text("Some fixed text")

    # Assert
    assert len(wrapper.doc.paragraphs) == 1
    assert wrapper.doc.paragraphs[0].text == "Some fixed text"
    assert wrapper.doc.paragraphs[0].runs[0].italic == True
    assert wrapper.doc.paragraphs[0].runs[0].font.size == Pt(12)
    assert wrapper.doc.paragraphs[0].runs[0].font.name == italic_font_name

