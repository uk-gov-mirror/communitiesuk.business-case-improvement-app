from unittest.mock import MagicMock
from apps.word_doc_services.create_document_objects import (
    create_table,
    TableDefinition
)

from docx import Document
 
def test_create_table_method_adds_table_to_document():
    # arrange
    test_doc = Document()

    # act
    create_table(TableDefinition.SINGLE_CELL_TEXT_BOX, test_doc)

    # assert
    assert len(test_doc.tables) == 1


def test_create_table_method_handles_invalid_table_definition():
    # arrange
    mock_enum = MagicMock(spec=TableDefinition)
    mock_enum.__int__.return_value = 999
    mock_enum.value = 999
    mock_enum.name = "TEST_ENUM"

    test_doc = Document()

    # act
    create_table(mock_enum, test_doc)

    # assert
    assert len(test_doc.tables) == 0


def test_create_table_method_adds_table_footer():
    # arrange
    test_doc_table_footer = "Test Table Footer"
    test_doc = Document()
    
    # act
    create_table(TableDefinition.SINGLE_CELL_TEXT_BOX, test_doc, test_doc_table_footer)
    test_para = test_doc.paragraphs[0]
    paragraph_count = len(test_doc.paragraphs)
    
    # assert
    assert paragraph_count == 1
    assert test_para.text == test_doc_table_footer

def test_create_table_method_can_add_a_table_for_each_table_definition():
    # arrange
    test_doc = Document()
    tbl_count: int = 0

    # act 
    for tbl_def in TableDefinition:
        tbl_count +=1
        create_table(tbl_def, test_doc)

    # assert
    assert tbl_count == len(test_doc.tables)

