
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx import Document as WordDoc

from django.views.decorators.http import require_POST

import json
import logging

from ..triage.models import (
    BusinessCaseResponse,
    BusinessCaseResponseBlock,
    BusinessCaseResponseSection,
    BusinessCaseResponseSummary,
    BusinessCase,
    BusinessCaseTriageResponse
)

model_exception_string: str = "Exception occurred submitting summary data to the {} model.\nException: {}"

summary_key_author: str = "author"
summary_key_directorate: str = "directorate"
summary_key_sro_scs: str = "sro_scs"
summary_key_text: str = "summary"
summary_key_whole_life_cost: str = "whole_life_cost"

logger = logging.getLogger(__name__)

'''
Summary:
    Class for the section content. Keep the section heading along with a list of all the
    paragraphs and tables that are within that section.
'''
class _SectionContent():
    def __init__(self):
        self.section_header = ""
        self.content = []

    def add_section_header(self, header: str):
        self.section_header = header

    def add_item_to_content(self, item: str | dict[str, str]):
        self.content.append(item)


@require_POST
def trigger_parsing(request):
    doc = WordDoc("FullDoc.docx")
    result = parse_word_document(doc)
    return result  


'''
Summary:
    Take the summary and document data sections and submit them to the models.
Returns: 
    -
Params:
    summary_section: dict. Contains key value pair of summary data.
    document_data: list[_SectionContent]. Needs splitting and itterating over.
'''
def submit_data_to_models(summary_section: dict, document_data: list[_SectionContent]) -> tuple[bool, bool]:
    summary_data_successful: bool = True
    document_data_successful: bool = True
    triage_response_object = BusinessCaseTriageResponse.objects.create()
    business_case_object, _ = BusinessCase.objects.get_or_create(business_case_triage_response=triage_response_object)

    BusinessCaseResponse.objects.create(
        uploaded_by="DefaultTestUser",
        business_case_id=business_case_object
    )

    business_case_response_object = BusinessCaseResponse.objects.get(business_case_id=business_case_object)

    # submit summary data to the model
    try:
        submit_summary_data(business_case_response_object, summary_section)
    except Exception as ex:
        summary_data_successful = False
        logger.error(model_exception_string.format("BusinessCaseResponseSummary", ex.__str__))

    # submit Document data to the models required.
    for data in document_data:
        business_Case_response_section=BusinessCaseResponseSection.objects.create(
            business_case_response_id=business_case_response_object,
            header_text=data.section_header
        )

        order_of_block: int = 1
        for c in data.content:
            if isinstance(c, str):
                paragraph_bytes= c.encode('utf-8')

                try:
                    BusinessCaseResponseBlock.objects.create(
                        business_case_response_section_id=business_Case_response_section,
                        block_type="Paragraph",
                        block_number=order_of_block,
                        block_data=paragraph_bytes
                    )
                except Exception as ex:
                    logger.error(model_exception_string.format("BusinessCaseResponseBlock, type: str", ex.__str__))
                    document_data_successful = False

            elif isinstance(c, dict):
                try:
                    json_con_data_string = json.dumps(c, separators=(',', ':'))
                    dict_bytes = json_con_data_string.encode('utf-8')

                    BusinessCaseResponseBlock.objects.create(
                        business_case_response_section_id=business_Case_response_section,
                        block_type="Table",
                        block_number=order_of_block,
                        block_data=dict_bytes
                    )
                except Exception as ex:
                    logger.error(model_exception_string.format("BusinessCaseResponseBlock, type: dict", ex.__str__))
                    document_data_successful = False              

            order_of_block += 1

    return summary_data_successful, document_data_successful


'''
Summary;
    Submit summary section data to the model.
Returns:
    -
Params:
    business_case_response_object: The business case response we need to link via FK in summary table.
    summary_section: Dictionary containing the summary data.
'''
def submit_summary_data(business_case_response_object: BusinessCaseResponse, summary_section: dict[str, str]):
    BusinessCaseResponseSummary.objects.create(
            business_case_response_id = business_case_response_object,
            summary_text = summary_section.get(summary_key_text, "-"),
            whole_of_life_cost = summary_section.get(summary_key_whole_life_cost, "-"),
            directorate = summary_section.get(summary_key_directorate, "-"),
            sro_scs = summary_section.get(summary_key_sro_scs, "-"),
            author = summary_section.get(summary_key_author, "-")
        )


'''
Summary:
    Read through a Word Doc and extract the Tables and Paragraphs contained in it.
    Currently, formatting is ignored.
Returns:
     -
Params:
    doc: Word Document.
'''
def parse_word_document(doc: Document):
    document_sections: list[_SectionContent] = []
    temp_section: _SectionContent
    summary_section: dict[str, str] = {}

    for s in doc.sections:
        content_stream = iter(s.iter_inner_content())
        content = next(content_stream, None)

        while content is not None:
            temp_section = _SectionContent()
            if not content or not content.style:
                continue

            if isinstance(content, Paragraph) and content.style.name == "Heading 1":
                if (content_text := escaped_string(content.text)) != "":
                    temp_section.add_section_header(content_text)
                    content = next(content_stream, None)
                    
                while content is not None:
                    if content.style and isinstance(content, Paragraph):
                        if content.style.name == "Heading 1":
                            break # new section, end current section
                        else:
                            if (content_text := escaped_string(content.text)) != "":
                                temp_section.add_item_to_content(content_text) 
                    elif isinstance(content, Table):
                        temp_section.add_item_to_content(
                            _extract_data_from_doc_table(content)
                        )

                    content = next(content_stream, None)
                document_sections.append(temp_section)

    summary_section = _get_summary_data(doc)
    submit_data_to_models(summary_section, document_sections)


'''
Summary:
    Get the summary data
Returns:
    dict:
        key - Name of the summary section
        value - Value of the section as str
Params:
    doc: Word Document
'''
def _get_summary_data(doc: Document) -> dict[str, str]:
    summary_data: dict[str, str] = {}
    directorate: str
    sro_scs: str
    author: str

    author, sro_scs, directorate = _get_data_from_details_table(doc)
 
    summary_data[summary_key_author] = author
    summary_data[summary_key_sro_scs] = sro_scs
    summary_data[summary_key_directorate] = directorate
    summary_data[summary_key_text] = _get_summary(doc)
    summary_data[summary_key_whole_life_cost] = _get_whole_life_cost(doc)

    return summary_data


'''
Summary:
    Get the whole life cost
Returns:
    str containing the whole life cost
Params:
    doc: Word Document
'''
def _get_whole_life_cost(doc: Document) -> str:
    whole_life_cost: str = "-"

    if doc and doc.tables:
        for tbl in doc.tables:
            if escaped_string(tbl.rows[0].cells[0].text) == "Whole Life Cost":
                whole_life_cost = f'£{escaped_string(tbl.rows[0].cells[1].text)}'
                break

    return whole_life_cost


'''
Summary:
    Get the Summary text.
Returns:
    str containing the Summary text
Params:
    doc: Word Document
'''
def _get_summary(doc: Document) -> str:
    summary:str = "-"

    if doc and doc.tables[2]:
        tbl = doc.tables[2]
        summary = tbl.rows[0].cells[0].text

    return summary


'''
Summary:
    Get the data in the details table from the document that are needed for the Summary section.
Returns:
    Tuple:
        Author name
        SRO or SCS name (approver)
        Directorate
Params:
    doc: Word Document.
'''
def _get_data_from_details_table(doc: Document) -> tuple[str, str, str]:
    # These should match the word doc content. If that value changes we should change here (maybe make a global string).
    # If changed by a user this search will fail.
    author_name_header: str = "author name:"
    sro_or_scs_header: str = "sro or area scs name (approver):"
    directorate_header: str = "directorate:"

    # default values per design
    author_name_result: str = "-"
    sro_or_scs_result: str = "-"
    directorate_result: str = "-"

    known_headers = {author_name_header, sro_or_scs_header , directorate_header}

    if doc and doc.tables[0]:
        tbl = doc.tables[0]

        for r in tbl.rows:
            cell_header: str = escaped_string(r.cells[0].text).lower()
            if cell_header in known_headers:
                summary_content = escaped_string(r.cells[1].text)
                match cell_header:
                    case header if header == author_name_header:
                        author_name_result = summary_content
                    case header if header == sro_or_scs_header:
                        sro_or_scs_result = summary_content
                    case header if header == directorate_header:
                        directorate_result = summary_content
                    case _:
                        continue

    return author_name_result, sro_or_scs_result, directorate_result


'''
Summary:
    Extract all the data from a table from a Word Doc into a single Dictionary.
Returns:
    dict:
        key - row index (determined by the Word Doc's row index for the table)
        value - a list of strings, representing the cell values from columns left->right
Params:
    tbl: the table from which to extract the data
'''
def _extract_data_from_doc_table(tbl: Table) -> dict[str, str]:
    tbl_dict: dict = {}

    for r in tbl.rows:
        tbl_dict.setdefault(r._index, [])
        for c in r.cells:
            tbl_dict[r._index].append(escaped_string(c.text))
    
    return tbl_dict


'''
Summary:
    Remove and format string in the way we need.
Returns:
    A formatted string. This exists beause when the template was pulled from the web into word, it
    came with nbs marks, and other markup as string that translated to hard text when saved.
    Also allows control over other elements like where the '£' sign goes in a string or remove double spaces, etc.
Params:
    value: any str.
'''
def escaped_string(value: str) -> str:
    return (value.replace("\xa0", " ") # non-breaking space
                .replace("\n", " ") # newline
                .replace("  ", " ") # replace any double spaces, or any created by the \n replace earlier
                .replace("£", "") # easier to strip this than check and replace if not etc. Then calling code can format.
                .strip()
            )
