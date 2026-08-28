from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.text.run import Run

import docx.oxml.ns
import docx.opc.constants
import logging

_mhclg_green_hex: str = "#00625E"
_blue_help_box_background_hex: str = "#E4F2FF"
_regular_font_colour_hex: str = "#000000"
_help_box_footer_font_colour_hex: str = "#E8E8E8"
_hyperlink_font_colour_hex: str = "#1D70B8"

bold_font_name: str = "Arial Bold"
regular_font_name: str = "Arial Regular"
italic_font_name: str = "Arial Italic"

logger = logging.getLogger(__name__)

def get_mhclg_green_rgb() -> RGBColor:
    return translate_hex_to_rgb(_mhclg_green_hex)

def get_blue_help_box_background_rgb() -> RGBColor:
    return translate_hex_to_rgb(_blue_help_box_background_hex)

def get_general_font_colour_rgb() -> RGBColor:
    return translate_hex_to_rgb(_regular_font_colour_hex)

def get_help_box_footer_font_colour_rgb() -> RGBColor:
    return translate_hex_to_rgb(_help_box_footer_font_colour_hex)

def translate_hex_to_rgb(hex_colour: str) -> RGBColor:
    return RGBColor(
        int(hex_colour[1:3],16),
        int(hex_colour[3:5],16),
        int(hex_colour[5:7],16)
    )


'''
Summary:
    Wrapper around the Word Document logic.
    
    Because Python-Docx uses proxy classes to handle the typing and what is returned, you can't have 
    objects like Paragraphs and Tables truly exist outside of the method to add_XXX() which returns the object.
    This requires having a doc somewhere which would couple things together.
    BusinessCaseWordDocument is wrapping the logic so as to separate calling code fromm logic as much as possible.
    
    Hopefully in this way the wrapper is souly responsible for styling/ handling the word document logic,
    and the methods are simply called by whatever needs a document created.
'''
class BusinessCaseWordDocumentWrapper:
    
    def __init__(self):
        self.doc = Document()


    def add_h1_section_header(self, header_text: str):
        p = self.get_basic_paragraph()
        r = p.add_run(header_text)
        r.font.color.rgb = get_mhclg_green_rgb()
        r.font.size = Pt(28)
        r.font.bold = True
        r.font.italic = False
        r.font.name = bold_font_name


    def add_h2_section_header(self, header_text: str):
        p = self.get_basic_paragraph()
        r = p.add_run(header_text)
        r.font.color.rgb = get_mhclg_green_rgb()
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.italic = False
        r.font.name = bold_font_name

    
    def add_paragraph(self, paragraph_content: str):
        p = self.get_basic_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        r = p.add_run(paragraph_content)
        self.style_paragraph_run(r)

    '''
    Summary:
        Get a paragraph with some basic styling applied that
        applies to all paragraphs.
    '''
    def get_basic_paragraph(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return p

    '''
    Summary:
        Style a paragraph run.
        Separate styling here so multiple methods can call it.
        This is generic styling for a regualr paragraph, i.e not a header.
    '''
    def style_paragraph_run(self, r: Run):
        r.font.color.rgb = get_general_font_colour_rgb()
        r.font.size = Pt(12)
        r.font.bold = False
        r.font.italic = False
        r.font.name = regular_font_name


    def add_bullet_point_list(self, items: list[str]):
        for item in items:
            p = self.doc.add_paragraph()
            p.style="List Bullet"
            p.paragraph_format.left_indent = Pt(45)
            r = p.add_run(item)
            r.font.color.rgb = get_general_font_colour_rgb()
            r.font.size = Pt(12)
            r.font.bold = False
            r.font.italic = False
            r.font.name = regular_font_name

    '''
    Summary:
        Add fixed text (e.g. text from triage we know about) to the document.
    '''
    def add_fixed_text(self, fixed_text: str):
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        r = p.add_run(fixed_text)
        r.font.color.rgb = get_general_font_colour_rgb()
        r.font.size = Pt(12)
        r.font.italic = True
        r.font.bold = False
        r.font.name = italic_font_name


    '''
    Summary:
        Add a Hyperlink to the Word Document.
        Because there is no explicit hyperlink class in Docx we have to make it using the Oxml.
        Take in a string that represents the URL, the entire text to display, and the text that should become the hyperlink.
    '''
    def add_hyperlink(self, url: str, text: str, text_to_replace_with_hyperlink):
        if (text_to_replace_with_hyperlink not in text):
            logger.warning(f"Text to replace with a hyperlink does not exist in the paragraph string. Paragraph: {text}, Text to replace: {text_to_replace_with_hyperlink}")
            self.add_paragraph(text)
            return
        
        p = self.doc.add_paragraph()

        part = p.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.ns.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = OxmlElement('w:r')

        new_run.append(self.get_hyperlink_run_style())
        new_run.text = text_to_replace_with_hyperlink
       
        sections = text.split(text_to_replace_with_hyperlink, 1)
        start = sections[0]
        end = sections[1]
        
        hyperlink.append(new_run)

        # join the text before the hyperlink, then the hyperlink, then the text after it
        start_run = p.add_run(f"{start.strip()} ")
        self.style_paragraph_run(start_run)

        p._p.append(hyperlink)

        end_run = p.add_run(f" {end.strip()}")
        self.style_paragraph_run(end_run)


    '''
    Summary:
        Style the hyperlink. Only the hyperlink itself, not the surrounding text.
    '''
    def get_hyperlink_run_style(self) -> BaseOxmlElement:
        # Create a new w:rPr element - Run Properties
        rPr = OxmlElement('w:rPr')
        
        color = OxmlElement('w:color')
        color.set(docx.oxml.ns.qn('w:val'), _hyperlink_font_colour_hex)
        rPr.append(color)

        underline = OxmlElement('w:u')
        underline.set(docx.oxml.ns.qn('w:val'), 'single')
        rPr.append(underline)

        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), bold_font_name)
        rFonts.set(qn('w:hAnsi'), bold_font_name)
        rPr.append(rFonts)

        # add Bold text
        b = OxmlElement('w:b')
        rPr.append(b)

        return rPr

    '''
    Summary;
        Save the document to the location required.
        If errors are encountered return a failing result.
    '''
    def save_document(self, save_location: str) -> bool:
        try:
            self.doc.save(save_location)
            return True
        except Exception as ex:
            logger.error(f"Error while saving Business Case template. Message: {ex.__str__}")
            return False

