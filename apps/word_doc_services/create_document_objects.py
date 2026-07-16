from docx.document import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE

from .table_definitions import TABLE_REGISTRY, TableDefinition
from .formatted_sections import add_text_with_default_formatting, aptos_font_name

def create_table(table_type: TableDefinition, doc: Document, table_footer: str = ""):
    table_data = next((item for item in TABLE_REGISTRY if item.definition == table_type), None)

    if not table_data or not table_data.cells:
        print(f"Table data not found. Skipping. table_type: {table_type.name}")
        return

    row_count = max(cell.row for cell in table_data.cells)
    column_count = max(cell.column for cell in table_data.cells)

    tbl = doc.add_table(row_count, column_count)
    tbl.style = "Table Grid" # adds gridlines to the table

    # resize the table column widths if they have been provided
    if table_data.column_widths:
        for col_idx, width in enumerate(table_data.column_widths):
            if col_idx < len(tbl.columns):
                for cell in tbl.columns[col_idx].cells:
                    cell.width = width

    # resize the table row heights if they have been provided
    if table_data.row_heights:
        for row_idx, height in enumerate(table_data.row_heights):
            if row_idx < len(tbl.rows):
                row = tbl.rows[row_idx]
                row.height = height
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # populate each cell ith the corresponding data, including some default formatting
    for cell_data in table_data.cells:
        r_idx = cell_data.row - 1
        c_idx = cell_data.column - 1

        tbl_cell = tbl.cell(r_idx, c_idx)
        if table_type is not TableDefinition.SINGLE_CELL_TEXT_BOX:
            tbl_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        for p_idx, p_data in enumerate(cell_data.paragraphs):
            para = tbl_cell.paragraphs[0] if p_idx == 0 else tbl_cell.add_paragraph()

            for run_data in p_data.runs:
                run = para.add_run(run_data.text)
                run.italic = run_data.italic
                run.bold = run_data.bold
                run.font.name = aptos_font_name
                run.font.size = Pt(12)

    if len(table_footer) > 0:
        tbl_footer_paragraph = doc.add_paragraph()
        tbl_footer_paragraph.paragraph_format.space_before = 0
        add_text_with_default_formatting(tbl_footer_paragraph, table_footer)

